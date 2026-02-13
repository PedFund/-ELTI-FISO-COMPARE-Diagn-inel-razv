import matplotlib
matplotlib.use('Agg') # Для сервера
import matplotlib.pyplot as plt
from flask import Flask, request, send_file, jsonify
import pandas as pd
import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from api.utils import process_dataframe

app = Flask(__name__)

def generate_chart(df, metric, title):
    """Рисует график сравнения"""
    levels = ["ниже нормативного", "нормативный", "выше нормативного"]
    
    # Считаем доли
    def get_pct(col):
        counts = df[col].value_counts().reindex(levels, fill_value=0)
        return (counts / len(df) * 100).fillna(0)

    pre = get_pct(f"{metric}_уровень_Start")
    post = get_pct(f"{metric}_уровень_End")

    fig, ax = plt.subplots(figsize=(6.5, 4))
    x = range(len(levels))
    width = 0.35

    # Цвета как в Excel (синий и голубой)
    bars1 = ax.bar([i - width/2 for i in x], pre, width, label='Начало года', color='#a6cee3', edgecolor='white')
    bars2 = ax.bar([i + width/2 for i in x], post, width, label='Конец года', color='#1f78b4', edgecolor='white')

    ax.set_ylabel('Доля детей (%)')
    ax.set_title(title, pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(["Ниже нормы", "Норма", "Выше нормы"])
    ax.legend()
    ax.grid(axis='y', linestyle='--', alpha=0.3)
    ax.set_ylim(0, 105)

    # Подписи значений
    for bars in [bars1, bars2]:
        for bar in bars:
            h = bar.get_height()
            if h > 0:
                ax.text(bar.get_x() + bar.get_width()/2, h + 1, f'{h:.1f}%', 
                        ha='center', va='bottom', fontsize=9)

    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150)
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

@app.route('/api/compare', methods=['POST'])
def compare():
    try:
        f1 = request.files.get('file_start')
        f2 = request.files.get('file_end')
        if not f1 or not f2: return jsonify({'error': 'Нужны оба файла'}), 400

        # Читаем и считаем метрики
        df1 = process_dataframe(f1)
        df2 = process_dataframe(f2)

        # Мержим по Коду
        df_merged = pd.merge(df1, df2, on="Код", suffixes=('_Start', '_End'), how='inner')
        if len(df_merged) == 0:
            return jsonify({'error': 'Нет совпадений по кодам детей. Проверьте колонку "Код ребёнка"'}), 400

        # --- WORD ---
        doc = Document()
        
        # Стиль заголовка
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        h = doc.add_heading('Сравнительный аналитический отчет', 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Выборка: {len(df_merged)} детей (прошедших обе диагностики).')
        doc.add_paragraph('Сравнение результатов начальной и итоговой диагностики.')

        metrics = [
            ("Когнитивное развитие", "Когнитивное развитие"),
            ("Воображение_итог", "Воображение"),
            ("ЭмСоцИнтеллект", "Эмоционально-социальный интеллект")
        ]

        for col, name in metrics:
            if f"{col}_уровень_Start" in df_merged.columns:
                doc.add_heading(f'Показатель: {name}', level=1)
                
                # Вставка графика
                img = generate_chart(df_merged, col, name)
                doc.add_picture(img, width=Inches(6))
                
                # Авто-вывод
                diff = (df_merged[f"{col}_уровень_End"] == "выше нормативного").mean() - \
                       (df_merged[f"{col}_уровень_Start"] == "выше нормативного").mean()
                
                p = doc.add_paragraph()
                if diff > 0:
                    runner = p.add_run(f"Доля детей с высоким уровнем выросла на {diff*100:.1f}%.")
                    runner.font.color.rgb = RGBColor(0, 100, 0) # Зеленый
                elif diff < 0:
                    p.add_run(f"Доля детей с высоким уровнем снизилась на {abs(diff)*100:.1f}%.")
                else:
                    p.add_run("Изменений в группе с высоким уровнем не зафиксировано.")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        # Формируем имя файла
        match = re.match(r'(\d+)-(\d+)', f1.filename)
        prefix = f"{match.group(1)}-{match.group(2)}" if match else "Report"
        filename = f"{prefix}_comparison.docx"

        res = send_file(output, as_attachment=True, download_name=filename)
        res.headers['X-Filename'] = filename
        res.headers['Access-Control-Expose-Headers'] = 'X-Filename'
        return res

    except Exception as e:
        return jsonify({'error': str(e)}), 500

app = app
