import matplotlib
matplotlib.use('Agg') # Для сервера
import matplotlib.pyplot as plt
from flask import Flask, request, send_file, jsonify
import pandas as pd
import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from api.utils import process_dataframe

app = Flask(__name__)

def generate_chart(df_start, df_end, metric, title):
    """Рисует график сравнения двух независимых групп (Срез А vs Срез Б)"""
    levels = ["ниже нормативного", "нормативный", "выше нормативного"]
    
    # Функция для расчета долей (в %)
    def get_pct(df, col_name):
        if col_name not in df.columns:
            return pd.Series([0, 0, 0], index=levels)
        counts = df[col_name].value_counts().reindex(levels, fill_value=0)
        total = len(df)
        if total == 0: return counts
        return (counts / total * 100).fillna(0)

    # Считаем проценты независимо
    pre = get_pct(df_start, f"{metric}_уровень")
    post = get_pct(df_end, f"{metric}_уровень")

    fig, ax = plt.subplots(figsize=(6.5, 4))
    x = range(len(levels))
    width = 0.35

    # Рисуем столбики (БЕЗ КОЛИЧЕСТВА В ЛЕГЕНДЕ)
    # Исправлено: убрал f"... ({len} чел.)" из label
    bars1 = ax.bar([i - width/2 for i in x], pre, width, label='Начало года', color='#a6cee3', edgecolor='white')
    bars2 = ax.bar([i + width/2 for i in x], post, width, label='Конец года', color='#1f78b4', edgecolor='white')

    ax.set_ylabel('Доля детей (%)')
    ax.set_title(title, pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(["Ниже нормы", "Норма", "Выше нормы"])
    ax.legend()
    ax.grid(axis='y', linestyle='--', alpha=0.3)
    ax.set_ylim(0, 105)

    # Подписи значений
    for bars in [bars1, bars2]:
        for bar in bars:
            h = bar.get_height()
            if h > 0:
                ax.text(bar.get_x() + bar.get_width()/2, h + 1, f'{h:.1f}%', 
                        ha='center', va='bottom', fontsize=9)

    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=150)
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

@app.route('/api/compare', methods=['POST'])
def compare():
    try:
        f1 = request.files.get('file_start')
        f2 = request.files.get('file_end')
        if not f1 or not f2: return jsonify({'error': 'Нужны оба файла'}), 400

        # Читаем файлы НЕЗАВИСИМО
        df_start = process_dataframe(f1)
        df_end = process_dataframe(f2)

        # --- ГЕНЕРАЦИЯ WORD ---
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        h = doc.add_heading('Сравнительный аналитический отчет', 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('Анализ проводится методом сравнения срезов (общие показатели группы на начало и конец периода).')
        
        # Информация о выборке в тексте отчета
        p = doc.add_paragraph()
        p.add_run(f'Выборка "Начало года": ').bold = True
        p.add_run(f'{len(df_start)} детей.')
        p = doc.add_paragraph()
        p.add_run(f'Выборка "Конец года": ').bold = True
        p.add_run(f'{len(df_end)} детей.')

        metrics = [
            ("Когнитивное развитие", "Когнитивное развитие"),
            ("Воображение_итог", "Воображение"),
            ("ЭмСоцИнтеллект", "Эмоционально-социальный интеллект")
        ]

        for col, name in metrics:
            if f"{col}_уровень" in df_start.columns and f"{col}_уровень" in df_end.columns:
                doc.add_heading(f'Показатель: {name}', level=1)
                
                img = generate_chart(df_start, df_end, col, name)
                doc.add_picture(img, width=Inches(6))
                
                # Авто-вывод
                pct_high_start = (df_start[f"{col}_уровень"] == "выше нормативного").mean()
                pct_high_end = (df_end[f"{col}_уровень"] == "выше нормативного").mean()
                
                diff = pct_high_end - pct_high_start
                
                p = doc.add_paragraph()
                if diff > 0:
                    runner = p.add_run(f"Доля детей с высоким уровнем выросла на {diff*100:.1f}%.")
                    runner.font.color.rgb = RGBColor(0, 100, 0) 
                elif diff < 0:
                    p.add_run(f"Доля детей с высоким уровнем снизилась на {abs(diff)*100:.1f}%.")
                else:
                    p.add_run("Изменений в группе с высоким уровнем не зафиксировано.")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        match = re.match(r'(\d+)-(\d+)', f1.filename)
        prefix = f"{match.group(1)}-{match.group(2)}" if match else "Report"
        filename = f"{prefix}_comparison_full_group.docx"

        res = send_file(output, as_attachment=True, download_name=filename)
        res.headers['X-Filename'] = filename
        res.headers['Access-Control-Expose-Headers'] = 'X-Filename'
        return res

    except Exception as e:
        return jsonify({'error': str(e)}), 500

app = app
